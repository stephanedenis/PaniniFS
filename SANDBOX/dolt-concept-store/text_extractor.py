#!/usr/bin/env python3
"""text_extractor.py — Unified text extraction for PaniniFS v4.0

Extracts structured text (paragraphs) from multiple document formats:
  PDF   → pdfminer.six
  EPUB  → ebooklib + BeautifulSoup
  DOCX  → python-docx
  HTML  → BeautifulSoup
  MD    → markdown-it-py
  TXT   → plain text with paragraph detection

Each extractor returns a list of ExtractedParagraph dataclasses.

v4.8: Unicode NFC normalization at the chokepoint (_clean_paragraphs)
      guarantees consistent keyword matching downstream.

Usage:
    from text_extractor import extract_document
    result = extract_document("/path/to/file.pdf")
    for para in result.paragraphs:
        print(para.text, para.index, para.section)
"""

import os
import re
import struct
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

# v4.8: Unicode NFC normalization — canonical form for consistent keyword matching
try:
    from text_normalizer import normalize_nfc
    HAS_NORMALIZER = True
except ImportError:
    import unicodedata
    def normalize_nfc(text: str) -> str:
        return unicodedata.normalize('NFC', text)
    HAS_NORMALIZER = False

# v4.7: Optional preamble normalizer for language-aware Gutenberg stripping
try:
    from gutenberg_preamble_normalizer import (
        classify_gutenberg_zones, detect_foreign_citations,
        strip_gutenberg_boilerplate, ZoneType,
    )
    HAS_PREAMBLE_NORMALIZER = True
except ImportError:
    HAS_PREAMBLE_NORMALIZER = False


# ─────────────────────────────────────────────────────────────────────────────
# DATA STRUCTURES
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class ExtractedParagraph:
    """A single paragraph extracted from a document."""
    text: str                          # cleaned paragraph text
    index: int                         # 0-based paragraph index
    section: Optional[str] = None      # chapter/section title if available
    page: Optional[int] = None         # page number (PDF only)
    word_count: int = 0                # word count
    char_count: int = 0                # character count


@dataclass
class ExtractionResult:
    """Result of extracting text from a document."""
    source_path: str                   # absolute path to source file
    format: str                        # detected format (pdf/epub/docx/html/md/txt)
    title: Optional[str] = None        # document title if available
    paragraphs: list = field(default_factory=list)  # list of ExtractedParagraph
    metadata: dict = field(default_factory=dict)     # format-specific metadata
    errors: list = field(default_factory=list)        # any extraction errors

    @property
    def total_words(self):
        return sum(p.word_count for p in self.paragraphs)

    @property
    def total_paragraphs(self):
        return len(self.paragraphs)


# ─────────────────────────────────────────────────────────────────────────────
# FORMAT DETECTION
# ─────────────────────────────────────────────────────────────────────────────

def detect_format(filepath: str) -> str:
    """Detect document format using magic numbers + extension fallback.
    
    Returns one of: 'pdf', 'epub', 'docx', 'html', 'md', 'txt'
    """
    path = Path(filepath)
    ext = path.suffix.lower()

    # Magic number detection (first 8 bytes)
    try:
        with open(filepath, 'rb') as f:
            header = f.read(8)
    except (OSError, IOError):
        return 'txt'  # fallback

    # PDF: %PDF-
    if header[:5] == b'%PDF-':
        return 'pdf'

    # ZIP-based formats (EPUB, DOCX)
    if header[:4] == b'PK\x03\x04':
        # Distinguish EPUB from DOCX by inspecting ZIP contents
        try:
            with zipfile.ZipFile(filepath, 'r') as zf:
                names = zf.namelist()
                if 'META-INF/container.xml' in names or any(
                    n.endswith('.opf') for n in names
                ):
                    return 'epub'
                if 'word/document.xml' in names or '[Content_Types].xml' in names:
                    return 'docx'
        except zipfile.BadZipFile:
            pass
        # Extension fallback for ZIP
        if ext == '.epub':
            return 'epub'
        if ext in ('.docx', '.doc'):
            return 'docx'

    # HTML detection (text-based)
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            first_1k = f.read(1000).lower()
        if '<html' in first_1k or '<!doctype html' in first_1k:
            return 'html'
    except Exception:
        pass

    # Extension-based fallback
    ext_map = {
        '.pdf': 'pdf', '.epub': 'epub', '.docx': 'docx',
        '.html': 'html', '.htm': 'html', '.xhtml': 'html',
        '.md': 'md', '.markdown': 'md', '.mkd': 'md',
        '.txt': 'txt', '.text': 'txt', '.rst': 'txt',
    }
    return ext_map.get(ext, 'txt')


# ─────────────────────────────────────────────────────────────────────────────
# PARAGRAPH CLEANING (shared)
# ─────────────────────────────────────────────────────────────────────────────

def _clean_paragraphs(raw_text: str, min_length: int = 10) -> list[str]:
    """Split raw text into cleaned paragraphs.
    
    Splits on double newlines, normalizes whitespace within paragraphs,
    and filters out very short fragments.
    v4.8: NFC normalization ensures canonical Unicode forms (é = U+00E9,
    not e + U+0301) so downstream keyword matching is consistent.
    """
    # Normalize line endings
    text = raw_text.replace('\r\n', '\n').replace('\r', '\n')

    # v4.8: Unicode NFC normalization — MUST come before any text comparison
    text = normalize_nfc(text)

    # Split on double newlines (paragraph boundaries)
    raw_paras = re.split(r'\n\s*\n', text)

    paragraphs = []
    for p in raw_paras:
        # Normalize internal whitespace (preserve single newlines as spaces)
        p = re.sub(r'[ \t]+', ' ', p)
        p = re.sub(r'\n', ' ', p)
        p = p.strip()
        if p and len(p) >= min_length:
            paragraphs.append(p)

    return paragraphs


def _make_extracted_paragraphs(
    texts: list[str],
    section: Optional[str] = None,
    page: Optional[int] = None,
    start_index: int = 0,
) -> list[ExtractedParagraph]:
    """Convert raw text strings to ExtractedParagraph objects."""
    result = []
    for i, text in enumerate(texts):
        result.append(ExtractedParagraph(
            text=text,
            index=start_index + i,
            section=section,
            page=page,
            word_count=len(text.split()),
            char_count=len(text),
        ))
    return result


# ─────────────────────────────────────────────────────────────────────────────
# PDF EXTRACTOR
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf(filepath: str) -> ExtractionResult:
    """Extract text from PDF using pdfminer.six."""
    from pdfminer.high_level import extract_text, extract_pages
    from pdfminer.layout import LTTextContainer, LTPage

    result = ExtractionResult(source_path=filepath, format='pdf')
    paragraphs = []
    para_idx = 0

    try:
        # Extract text page by page for page-number tracking
        for page_num, page_layout in enumerate(extract_pages(filepath), 1):
            page_text = []
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    text = element.get_text().strip()
                    if text:
                        page_text.append(text)

            if page_text:
                combined = '\n\n'.join(page_text)
                cleaned = _clean_paragraphs(combined)
                new_paras = _make_extracted_paragraphs(
                    cleaned, page=page_num, start_index=para_idx
                )
                paragraphs.extend(new_paras)
                para_idx += len(new_paras)

        result.paragraphs = paragraphs
        result.metadata['total_pages'] = page_num if paragraphs else 0

        # Try to extract title from first page text
        if paragraphs:
            first_para = paragraphs[0].text
            if len(first_para) < 200:  # Short first paragraph = likely title
                result.title = first_para

    except Exception as e:
        result.errors.append(f"PDF extraction error: {e}")
        # Fallback: try simple full-text extraction
        try:
            full_text = extract_text(filepath)
            cleaned = _clean_paragraphs(full_text)
            result.paragraphs = _make_extracted_paragraphs(cleaned)
        except Exception as e2:
            result.errors.append(f"PDF fallback error: {e2}")

    return result


# ─────────────────────────────────────────────────────────────────────────────
# EPUB EXTRACTOR
# ─────────────────────────────────────────────────────────────────────────────

def _extract_epub(filepath: str) -> ExtractionResult:
    """Extract text from EPUB using ebooklib + BeautifulSoup."""
    from ebooklib import epub
    from bs4 import BeautifulSoup

    result = ExtractionResult(source_path=filepath, format='epub')
    paragraphs = []
    para_idx = 0

    try:
        book = epub.read_epub(filepath, options={"ignore_ncx": True})

        # Extract title from metadata
        title_items = book.get_metadata('DC', 'title')
        if title_items:
            result.title = title_items[0][0]

        # Extract author(s)
        authors = book.get_metadata('DC', 'creator')
        if authors:
            result.metadata['authors'] = [a[0] for a in authors]

        # Extract language
        langs = book.get_metadata('DC', 'language')
        if langs:
            result.metadata['language'] = langs[0][0]

        # Process spine items (reading order)
        for item in book.get_items_of_type(9):  # ITEM_DOCUMENT = 9
            content = item.get_content().decode('utf-8', errors='replace')
            soup = BeautifulSoup(content, 'html.parser')

            # Try to detect chapter title
            chapter_title = None
            for tag in ('h1', 'h2', 'h3'):
                heading = soup.find(tag)
                if heading:
                    chapter_title = heading.get_text(strip=True)
                    break

            # Extract paragraphs from <p> tags
            p_tags = soup.find_all('p')
            if p_tags:
                for p_tag in p_tags:
                    text = p_tag.get_text(strip=True)
                    if text and len(text) >= 10:
                        paragraphs.append(ExtractedParagraph(
                            text=text,
                            index=para_idx,
                            section=chapter_title,
                            word_count=len(text.split()),
                            char_count=len(text),
                        ))
                        para_idx += 1
            else:
                # Fallback: extract all text from body
                body = soup.find('body')
                if body:
                    raw = body.get_text(separator='\n\n')
                    cleaned = _clean_paragraphs(raw)
                    new_paras = _make_extracted_paragraphs(
                        cleaned, section=chapter_title, start_index=para_idx
                    )
                    paragraphs.extend(new_paras)
                    para_idx += len(new_paras)

        result.paragraphs = paragraphs

    except Exception as e:
        result.errors.append(f"EPUB extraction error: {e}")

    return result


# ─────────────────────────────────────────────────────────────────────────────
# DOCX EXTRACTOR
# ─────────────────────────────────────────────────────────────────────────────

def _extract_docx(filepath: str) -> ExtractionResult:
    """Extract text from DOCX using python-docx."""
    from docx import Document

    result = ExtractionResult(source_path=filepath, format='docx')
    paragraphs = []
    para_idx = 0
    current_section = None

    try:
        doc = Document(filepath)

        # Extract core properties
        props = doc.core_properties
        if props.title:
            result.title = props.title
        if props.author:
            result.metadata['author'] = props.author

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings
            if para.style and para.style.name.startswith('Heading'):
                current_section = text
                continue

            if len(text) >= 10:
                paragraphs.append(ExtractedParagraph(
                    text=text,
                    index=para_idx,
                    section=current_section,
                    word_count=len(text.split()),
                    char_count=len(text),
                ))
                para_idx += 1

        result.paragraphs = paragraphs

    except Exception as e:
        result.errors.append(f"DOCX extraction error: {e}")

    return result


# ─────────────────────────────────────────────────────────────────────────────
# HTML EXTRACTOR
# ─────────────────────────────────────────────────────────────────────────────

def _extract_html(filepath: str) -> ExtractionResult:
    """Extract text from HTML using BeautifulSoup."""
    from bs4 import BeautifulSoup

    result = ExtractionResult(source_path=filepath, format='html')
    paragraphs = []
    para_idx = 0

    try:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()

        soup = BeautifulSoup(content, 'html.parser')

        # Extract title
        title_tag = soup.find('title')
        if title_tag:
            result.title = title_tag.get_text(strip=True)

        # Extract meta information
        for meta in soup.find_all('meta'):
            name = meta.get('name', '').lower()
            content_val = meta.get('content', '')
            if name in ('author', 'description', 'language'):
                result.metadata[name] = content_val

        # Remove script/style elements
        for tag in soup(['script', 'style', 'nav', 'header', 'footer']):
            tag.decompose()

        # Try article/main content first, then body
        content_root = (
            soup.find('article') or
            soup.find('main') or
            soup.find('body') or
            soup
        )

        # Track current section from headings
        current_section = None
        for element in content_root.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'div', 'blockquote']):
            if element.name in ('h1', 'h2', 'h3', 'h4'):
                current_section = element.get_text(strip=True)
                continue

            text = element.get_text(strip=True)
            if text and len(text) >= 10:
                # Avoid duplicates from nested elements
                if paragraphs and text == paragraphs[-1].text:
                    continue
                paragraphs.append(ExtractedParagraph(
                    text=text,
                    index=para_idx,
                    section=current_section,
                    word_count=len(text.split()),
                    char_count=len(text),
                ))
                para_idx += 1

        result.paragraphs = paragraphs

    except Exception as e:
        result.errors.append(f"HTML extraction error: {e}")

    return result


# ─────────────────────────────────────────────────────────────────────────────
# MARKDOWN EXTRACTOR
# ─────────────────────────────────────────────────────────────────────────────

def _extract_markdown(filepath: str) -> ExtractionResult:
    """Extract text from Markdown using markdown-it-py."""
    from markdown_it import MarkdownIt

    result = ExtractionResult(source_path=filepath, format='md')
    paragraphs = []
    para_idx = 0

    try:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()

        md = MarkdownIt()
        tokens = md.parse(content)

        current_section = None
        i = 0
        while i < len(tokens):
            token = tokens[i]

            # Track heading sections
            if token.type == 'heading_open':
                # Next token should be heading content (inline)
                if i + 1 < len(tokens) and tokens[i + 1].type == 'inline':
                    current_section = tokens[i + 1].content
                    i += 3  # skip heading_open, inline, heading_close
                    continue

            # Extract paragraph content
            if token.type == 'paragraph_open':
                if i + 1 < len(tokens) and tokens[i + 1].type == 'inline':
                    text = tokens[i + 1].content.strip()
                    if text and len(text) >= 10:
                        # Strip inline markdown formatting
                        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold
                        text = re.sub(r'\*(.+?)\*', r'\1', text)      # italic
                        text = re.sub(r'`(.+?)`', r'\1', text)        # code
                        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # links

                        paragraphs.append(ExtractedParagraph(
                            text=text,
                            index=para_idx,
                            section=current_section,
                            word_count=len(text.split()),
                            char_count=len(text),
                        ))
                        para_idx += 1

            i += 1

        result.paragraphs = paragraphs

        # Title = first heading
        if current_section is None:
            for token in tokens:
                if token.type == 'inline' and tokens[tokens.index(token) - 1].type == 'heading_open':
                    result.title = token.content
                    break
        else:
            # Re-scan for first H1
            for j, t in enumerate(tokens):
                if t.type == 'heading_open' and t.tag == 'h1':
                    if j + 1 < len(tokens):
                        result.title = tokens[j + 1].content
                    break

    except Exception as e:
        result.errors.append(f"Markdown extraction error: {e}")

    return result


# ─────────────────────────────────────────────────────────────────────────────
# PLAIN TEXT EXTRACTOR
# ─────────────────────────────────────────────────────────────────────────────

def _extract_txt(filepath: str) -> ExtractionResult:
    """Extract text from plain text files."""
    import chardet

    result = ExtractionResult(source_path=filepath, format='txt')

    try:
        # Detect encoding
        with open(filepath, 'rb') as f:
            raw = f.read()
        detected = chardet.detect(raw)
        encoding = detected.get('encoding', 'utf-8') or 'utf-8'
        result.metadata['encoding'] = encoding
        result.metadata['encoding_confidence'] = detected.get('confidence', 0)

        text = raw.decode(encoding, errors='replace')

        # v4.8: NFC normalize immediately after decode — repairs mojibake
        # from cp1252/latin-1 misdetection and ensures canonical forms
        text = normalize_nfc(text)

        # Strip Gutenberg-style headers/footers if present
        text_lower = text[:3000].lower()
        if 'project gutenberg' in text_lower or 'projet gutenberg' in text_lower:
            if HAS_PREAMBLE_NORMALIZER:
                # v4.7: Language-aware zone classification + boilerplate removal
                zones = classify_gutenberg_zones(text)
                text = strip_gutenberg_boilerplate(text)
                # Enrich metadata with zone and citation info
                result.metadata['gutenberg_zones'] = [
                    {
                        'type': z.zone_type.name.lower(),
                        'language': z.language,
                        'confidence': z.confidence,
                        'chars': z.end_char - z.start_char,
                        'semantic_id': z.metadata.get('semantic_id', ''),
                        'equivalent_across_languages': z.metadata.get(
                            'equivalent_across_languages', False
                        ),
                    }
                    for z in zones
                ]
                result.metadata['gutenberg_stripped'] = True
                result.metadata['gutenberg_normalizer_version'] = '4.7'
            else:
                # Legacy: English-only marker-based stripping
                start_markers = [
                    '*** START OF THE PROJECT GUTENBERG',
                    '*** START OF THIS PROJECT GUTENBERG',
                    '*END*THE SMALL PRINT',
                ]
                for marker in start_markers:
                    pos = text.upper().find(marker.upper())
                    if pos >= 0:
                        text = text[pos + len(marker):]
                        nl = text.find('\n')
                        if nl >= 0:
                            text = text[nl + 1:]
                        break

                end_markers = [
                    '*** END OF THE PROJECT GUTENBERG',
                    '*** END OF THIS PROJECT GUTENBERG',
                    'End of the Project Gutenberg',
                    'End of Project Gutenberg',
                ]
                for marker in end_markers:
                    pos = text.upper().find(marker.upper())
                    if pos >= 0:
                        text = text[:pos]
                        break

        cleaned = _clean_paragraphs(text)
        result.paragraphs = _make_extracted_paragraphs(cleaned)

        # Title heuristic: first short paragraph
        if result.paragraphs and result.paragraphs[0].word_count < 15:
            result.title = result.paragraphs[0].text

    except Exception as e:
        result.errors.append(f"TXT extraction error: {e}")

    return result


# ─────────────────────────────────────────────────────────────────────────────
# UNIFIED ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

_EXTRACTORS = {
    'pdf':  _extract_pdf,
    'epub': _extract_epub,
    'docx': _extract_docx,
    'html': _extract_html,
    'md':   _extract_markdown,
    'txt':  _extract_txt,
}


def extract_document(filepath: str, force_format: Optional[str] = None) -> ExtractionResult:
    """Extract structured text from any supported document format.
    
    Args:
        filepath: Path to the document file.
        force_format: Override format detection ('pdf', 'epub', 'docx', 'html', 'md', 'txt').
    
    Returns:
        ExtractionResult with paragraphs, metadata, and any errors.
    
    Example:
        >>> result = extract_document("my_book.epub")
        >>> print(f"{result.total_paragraphs} paragraphs, {result.total_words} words")
        >>> for p in result.paragraphs[:5]:
        ...     print(f"  [{p.section}] {p.text[:80]}...")
    """
    filepath = os.path.abspath(filepath)
    if not os.path.exists(filepath):
        return ExtractionResult(
            source_path=filepath,
            format='unknown',
            errors=[f"File not found: {filepath}"]
        )

    fmt = force_format or detect_format(filepath)
    extractor = _EXTRACTORS.get(fmt, _extract_txt)

    result = extractor(filepath)
    result.metadata['detected_format'] = fmt
    result.metadata['file_size_bytes'] = os.path.getsize(filepath)

    return result


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print("Usage: python text_extractor.py <file> [format]")
        print("Formats: pdf, epub, docx, html, md, txt (auto-detected if omitted)")
        sys.exit(1)

    filepath = sys.argv[1]
    force_fmt = sys.argv[2] if len(sys.argv) > 2 else None

    result = extract_document(filepath, force_format=force_fmt)

    print(f"\n{'=' * 70}")
    print(f"TEXT EXTRACTION: {os.path.basename(filepath)}")
    print(f"{'=' * 70}")
    print(f"  Format:      {result.format}")
    print(f"  Title:       {result.title or '(none)'}")
    print(f"  Paragraphs:  {result.total_paragraphs}")
    print(f"  Total words: {result.total_words}")
    if result.metadata:
        print(f"  Metadata:    {result.metadata}")
    if result.errors:
        print(f"  Errors:      {result.errors}")

    print(f"\n{'─' * 70}")
    print(f"First 5 paragraphs:")
    print(f"{'─' * 70}")
    for p in result.paragraphs[:5]:
        section = f"[{p.section}] " if p.section else ""
        page = f" (p.{p.page})" if p.page else ""
        print(f"  §{p.index}{page}: {section}{p.text[:120]}{'...' if len(p.text) > 120 else ''}")
        print(f"         ({p.word_count} mots)")
